import docx
from datetime import datetime, timedelta
import requests
import os
from pathlib import Path
from bs4 import BeautifulSoup as bs
from json import load
import subprocess
TXST_CALENDAR = None
OUT_PATH = None


# find class dates between start, end, only on weekdays
def build_dates(start, end, weekdays, holidays):
    '''
    returns a dictionary
    {
        "dates": [
            "Mon 01",
            etc...
        ],
        "topic": [
            "blank",
            etc...
        ],
        "holidays": [
            "MLK Day",
            etc...
        ]
    }
    '''
    week_to_day = {
        'Sunday': 'S',
        'Monday': 'M',
        'Tuesday': 'T',
        'Wednesday': 'W',
        'Thursday': 'R',
        'Friday': 'F',
        'Saturday': 'A',
    }
    date = start  # init date counter
    day = timedelta(days=1)  # incrementer for date
    calendar = {  # container for table data
        "dates": [],
        "topics": [],
        "holidays": [],
    }

    while date <= end:
        # only add a date if it's one of the user-submitted weekdays
        if week_to_day[date.strftime('%A')] in weekdays:
            holi = False  # true if I found a holiday

            # append this date to calendar['dates']
            calendar['dates'].append(date.strftime('%b %d'))

            # add 'final exams' on last day, or append placeholder for syllabus
            if date.date() == end.date():
                calendar['topics'].append("Final Exams")
            else:
                calendar['topics'].append('')

            # add holiday info if this date is a holiday
            for holiday in holidays:
                if holiday['start'] <= date <= holiday['end']:
                    calendar['holidays'].append(
                        f"Holiday - {holiday['name']}")
                    holi = True
                    break

            if not holi:  # append empty string if I didn't find a holiday
                calendar['holidays'].append('')  # add logic for holidays
        date += day  # increment date

    return calendar


# gets start(datetime(MM DD YY)) end(datetime(MM DD YY)) weekdays([uMTWhFS])
def get_input():
    weekdays = None
    start = None
    end = None
    format = None
    style = None

    # ask until I recieve a "" or "1" - output style
    while not style:
        style = input("Would you prefer a table output or a continuous "
                      "monthly calendar?\n"
                      "(press Enter for table or 1 for continuous)\n")
        if style == "1":
            style = "continuous"
        elif style == "":
            style = "table"
        else:
            print("Not a recognized style.")
            style = None

    # ask until I recieve good data - Class days
    while not weekdays:
        weekdays = input("\nWhich weekdays? Type a one letter code for each"
                         " day, with no spaces: [SMTWRFA]\n")
        for char in weekdays:
            if char not in "SMTWRFA":
                print("Bad day codes")
                weekdays = None

    # ask until I receive good data - first day of classes
    while not start:
        try:
            start = input("\nWhat is the first day of classes? MM DD YY\n")
            start = datetime.strptime(start, "%m %d %y")
        except ValueError:
            print("Bad Date")
            start = None

    # ask until I receive good data - final exam day
    while not end:
        try:
            end = input("\nWhat is the last day of class? MM DD YY\n")
            end = datetime.strptime(end, "%m %d %y")
        except ValueError:
            print("Bad Date")
            start = None

    # ask until I recieve something (only runs once) - format option
    while not format:
        format = input("\nWhat format would you like to use?\n"
                       "(press enter to accept no formatting)\n"
                       "(type 1 to use Light Shading format)\n")
        if format == "1":
            format = "Light Shading"
        elif format == "":
            format = "None"
    if format == "None":
        format = None

    return {
        "start": start,
        "end": end,
        "weekdays": weekdays,
        "format": format,
        "style": style
    }


# scrape TXST academic calendar for holidays between start and end
def get_TXST_holidays(start, end):
    # get html from TXST website
    global TXST_CALENDAR
    request = requests.get(TXST_CALENDAR)
    soup = bs(request.content, 'lxml')  # load html using Beautiful Soup

    holidays = []
    for td in soup.findAll('td'):
        # get all <td> if they contain a <div data-categories="Holidays">
        if td.find('div', attrs={'data-categories': 'Holidays'}):
            holiday = {}  # init a dict

            # pull name of holiday
            holiday['name'] = td.find(
                'div',
                attrs={'class': 'event-title'}
            ).text

            # pull start date
            holiday['start'] = datetime.strptime(td.find(
                'div',
                attrs={'class': 'event-data'}
            )['data-startdate'], "%Y-%m-%d")

            # pull end date
            holiday['end'] = datetime.strptime(td.find(
                'div',
                attrs={'class': 'event-data'}
            )['data-enddate'], "%Y-%m-%d")

            # append the details if the dates overlap with class dates
            if (start <= holiday['start'] <= end or
               start <= holiday['end'] <= end):
                holidays.append(holiday)

    return holidays


# take a 3 column table and a class calendar,
# fill in the table with the calendar info.
def build_table(table, calendar, format):
    # styling the table
    if format:
        table.columns[0].width = docx.shared.Inches(1).emu
        table.columns[1].width = docx.shared.Inches(2.5).emu
        table.columns[2].width = docx.shared.Inches(2.5).emu
        try:
            table.style = format
        except KeyError:
            print(f"There is no '{format}' format. Using default format.")

    # fill in the header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Date"
    hdr_cells[1].text = "Topics"
    hdr_cells[2].text = "Assignments"

    # for each entry in the list, append an entry to the table
    i = 0
    while i < len(calendar['dates']):
        row_cells = table.add_row().cells
        row_cells[0].text = calendar['dates'][i]
        row_cells[2].text = calendar['holidays'][i]
        i += 1


def load_config():
    try:
        file = open('files/config.json', 'r')
        config = load(file)
        file.close()
        global TXST_CALENDAR
        TXST_CALENDAR = config['TXST_CALENDAR']
        global OUT_PATH
        OUT_PATH = Path(config['OUT_PATH'])
    except Exception as e:
        print("Failed to load config.json")
        print(e)


def build_cont_calendar(table, calendar, start, end, format):
    # GREY = docx.shared.RGBColor(155, 155, 155)
    if format:
        try:
            table.style = format
        except KeyError:
            print(f"There is no '{format}' format. Using default format.")

    day = timedelta(days=1)
    week_begin = start
    week_end = end

    while week_begin.strftime("%a") != "Sun":
        week_begin -= day
    while week_end.strftime("%a") != "Sat":
        week_end += day

    # fill in the header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Sun"
    hdr_cells[1].text = "Mon"
    hdr_cells[2].text = "Tue"
    hdr_cells[3].text = "Wed"
    hdr_cells[4].text = "Thu"
    hdr_cells[5].text = "Fri"
    hdr_cells[6].text = "Sat"

    cal_day = week_begin
    print(calendar)

    while cal_day < week_end:
        # month_accent = int(cal_day.strftime("%m")) % 2
        row = table.add_row().cells
        for cell in row:
            if cal_day.strftime('%d') == "01":
                cell.text = cal_day.strftime('%B') + " "
            else:
                cell.text = ""

            if start.date() <= cal_day.date() <= end.date():
                cell.text += cal_day.strftime('%d')

                try:
                    index = calendar["dates"].index(cal_day.strftime('%b %d'))
                except ValueError:
                    index = -1
                if index > 0 and calendar["holidays"][index]:
                    cell.text += ("\n" + calendar["holidays"][index])
                elif index > 0:
                    cell.text += "\nClass Day"
                else:
                    cell.text += "\n" + cal_day.strftime('%b %d')
            cal_day += day


##############
# MAIN TREE ##
##############
if __name__ == "__main__":
    os.system('cls' if os.name == 'nt' else 'clear')
    load_config()
    print("\n                   ######################\n"
          "                   # Welcome to Calends #\n"
          "                   #      version 0.9   #\n"
          "                   ######################\n\n")

    # get table data
    inputs = get_input()  # get inputs
    # get observed holidays
    holidays = get_TXST_holidays(inputs["start"], inputs["end"])
    class_dates = build_dates(  # populate list
        inputs["start"],
        inputs["end"],
        inputs["weekdays"],
        holidays
    )

    # create .docx table
    document = docx.Document()  # init output doc

    if inputs["style"] == "continuous":  # convert list to calendar
        table = document.add_table(rows=1, cols=7)
        build_cont_calendar(
            table,
            class_dates,
            inputs["start"],
            inputs["end"],
            inputs["format"]
        )
    else:  # convert list to table
        table = document.add_table(rows=1, cols=3)
        table.autofit = False
        build_table(table, class_dates, inputs["format"])

    # save document to file.
    file_name = 'calends-output.docx'
    filepath = os.path.join(OUT_PATH, file_name)
    try:
        document.save(filepath)

        # open explorer to the new file and say goodbye
        subprocess.Popen(fr'explorer /select,"{str(filepath)}')
        print(f"\nPrinted calendar to {filepath}\nGoodbye!\n")
    except PermissionError:
        print(f"Failed to print to {filepath}")
